import os
from flask import Flask, render_template, send_from_directory, jsonify, request, send_file
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from datetime import datetime
import io


# Database setup
class Base(DeclarativeBase):
    pass

db = SQLAlchemy(model_class=Base)

# Create Flask app
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "a-secret-key-for-betting")

# Database configuration
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}
db.init_app(app)

# Models
class Game(db.Model):
    __tablename__ = 'games'
    
    id = db.Column(db.String(50), primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    home_team = db.Column(db.String(100), nullable=False)
    away_team = db.Column(db.String(100), nullable=False)
    home_odd = db.Column(db.Numeric(5, 2), nullable=False)
    draw_odd = db.Column(db.Numeric(5, 2), nullable=False)
    away_odd = db.Column(db.Numeric(5, 2), nullable=False)
    status = db.Column(db.String(20), default='active')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    def to_dict(self):
        return {
            'id': self.id,
            'name': self.name,
            'homeTeam': self.home_team,
            'awayTeam': self.away_team,
            'odds': {
                'home': float(self.home_odd),
                'draw': float(self.draw_odd),
                'away': float(self.away_odd)
            },
            'status': self.status,
            'createdAt': self.created_at.isoformat() if self.created_at else None
        }


class Bet(db.Model):
    __tablename__ = 'bets'
    
    id = db.Column(db.String(50), primary_key=True)
    player = db.Column(db.String(100), nullable=False)
    game_id = db.Column(db.String(50), db.ForeignKey('games.id'), nullable=False)
    game_name = db.Column(db.String(200), nullable=False)
    bet_type = db.Column(db.String(10), nullable=False)
    amount = db.Column(db.Numeric(12, 2), nullable=False)
    odd = db.Column(db.Numeric(5, 2), nullable=False)
    possible_win = db.Column(db.Numeric(12, 2), nullable=False)
    status = db.Column(db.String(20), default='pending')
    home_team = db.Column(db.String(100))
    away_team = db.Column(db.String(100))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    game = db.relationship('Game', backref='bets')
    
    def to_dict(self):
        return {
            'id': self.id,
            'player': self.player,
            'gameId': self.game_id,
            'gameName': self.game_name,
            'type': self.bet_type,
            'amount': float(self.amount),
            'odd': float(self.odd),
            'possibleWin': float(self.possible_win),
            'status': self.status,
            'createdAt': self.created_at.isoformat() if self.created_at else None,
            'gameDetails': {
                'homeTeam': self.home_team,
                'awayTeam': self.away_team
            }
        }

# Create tables
with app.app_context():
    db.create_all()

# Static routes
@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/style.css')
def style():
    return send_from_directory('.', 'style.css')

@app.route('/script.js')
def script():
    return send_from_directory('.', 'script.js')

@app.route('/manifest.json')
def manifest():
    return send_from_directory('.', 'manifest.json')

# API ROUTES - JOGOS

@app.route('/api/games', methods=['GET'])
def get_games():
    """Retorna todos os jogos ativos"""
    try:
        games = Game.query.filter_by(status='active').all()
        return jsonify({'games': [game.to_dict() for game in games]}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/games', methods=['POST'])
def create_game():
    """Cria um novo jogo"""
    try:
        data = request.get_json()
        
        new_game = Game(
            id=data['id'],
            name=data['name'],
            home_team=data['homeTeam'],
            away_team=data['awayTeam'],
            home_odd=data['odds']['home'],
            draw_odd=data['odds']['draw'],
            away_odd=data['odds']['away'],
            status=data.get('status', 'active')
        )
        
        db.session.add(new_game)
        db.session.commit()
        
        return jsonify({'game': new_game.to_dict()}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/games/<game_id>', methods=['DELETE'])
def delete_game(game_id):
    """Remove um jogo"""
    try:
        game = Game.query.get_or_404(game_id)
        db.session.delete(game)
        db.session.commit()
        
        return jsonify({'message': 'Jogo removido com sucesso'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

# API ROUTES - APOSTAS

@app.route('/api/bets', methods=['GET'])
def get_bets():
    """Retorna todas as apostas"""
    try:
        bets = Bet.query.all()
        return jsonify({'bets': [bet.to_dict() for bet in bets]}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/bets', methods=['POST'])
def create_bet():
    """Cria uma nova aposta"""
    try:
        data = request.get_json()
        
        new_bet = Bet(
            id=data['id'],
            player=data['player'],
            game_id=data['gameId'],
            game_name=data['gameName'],
            bet_type=data['type'],
            amount=data['amount'],
            odd=data['odd'],
            possible_win=data['possibleWin'],
            home_team=data['gameDetails']['homeTeam'],
            away_team=data['gameDetails']['awayTeam'],
            status=data.get('status', 'pending')
        )
        
        db.session.add(new_bet)
        db.session.commit()
        
        return jsonify({'bet': new_bet.to_dict()}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/bets/<bet_id>', methods=['PUT'])
def update_bet(bet_id):
    """Atualiza o status de uma aposta"""
    try:
        data = request.get_json()
        bet = Bet.query.get_or_404(bet_id)
        
        bet.status = data['status']
        db.session.commit()
        
        return jsonify({'bet': bet.to_dict()}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/bets/<bet_id>', methods=['DELETE'])
def delete_bet(bet_id):
    """Remove uma aposta"""
    try:
        bet = Bet.query.get_or_404(bet_id)
        db.session.delete(bet)
        db.session.commit()
        
        return jsonify({'message': 'Aposta removida com sucesso'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

# RELATÓRIOS

@app.route('/generate-pdf-report', methods=['POST'])
def generate_pdf_report():
    """Gera relatório PDF das apostas"""
    try:
        # Obter dados do banco de dados
        games = Game.query.all()
        bets = Bet.query.all()
        
        # Converter para dicionários 
        games_data = [game.to_dict() for game in games]
        bets_data = [bet.to_dict() for bet in bets]
        
        # Create PDF in memory
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        # Build content
        content = []
        
        # Title
        title = Paragraph("🏆 MASTER LEAGUE - Relatório de Apostas", title_style)
        content.append(title)
        
        # Date
        date_para = Paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal'])
        date_para.alignment = 2  # Right alignment
        content.append(date_para)
        content.append(Spacer(1, 30))
        
        # Statistics
        total_bets = len(bets_data)
        total_amount = sum(bet['amount'] for bet in bets_data)
        pending_bets = len([bet for bet in bets_data if bet['status'] == 'pending'])
        won_bets = len([bet for bet in bets_data if bet['status'] == 'won'])
        lost_bets = len([bet for bet in bets_data if bet['status'] == 'lost'])
        
        stats_data = [
            ['Estatística', 'Valor'],
            ['Total de Apostas', str(total_bets)],
            ['Valor Total Apostado', f"€{total_amount:,.2f}".replace(',', '.')],
            ['Apostas Pendentes', str(pending_bets)],
            ['Apostas Ganhas', str(won_bets)],
            ['Apostas Perdidas', str(lost_bets)]
        ]
        
        stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        content.append(stats_table)
        content.append(Spacer(1, 30))
        
        # Detailed bets table
        if bets_data:
            content.append(Paragraph("Detalhes das Apostas", styles['Heading2']))
            content.append(Spacer(1, 12))
            
            bets_table_data = [['Jogador', 'Jogo', 'Aposta', 'Valor (€)', 'Odd', 'Possível Ganho (€)', 'Status']]
            
            for bet in bets_data:
                bet_type_text = ''
                if bet['type'] == 'home':
                    bet_type_text = bet['gameDetails']['homeTeam']
                elif bet['type'] == 'away':
                    bet_type_text = bet['gameDetails']['awayTeam']
                else:
                    bet_type_text = 'Empate'
                
                status_text = {
                    'pending': 'Pendente',
                    'won': 'GANHOU ✓',
                    'lost': 'PERDEU ✗'
                }.get(bet['status'], 'Pendente')
                
                bets_table_data.append([
                    bet['player'],
                    bet['gameName'],
                    bet_type_text,
                    f"€{bet['amount']:,.2f}".replace(',', '.'),
                    str(bet['odd']),
                    f"€{bet['possibleWin']:,.2f}".replace(',', '.'),
                    status_text
                ])
            
            bets_table = Table(bets_table_data, colWidths=[1.2*inch, 1.2*inch, 1*inch, 0.8*inch, 0.6*inch, 1*inch, 0.8*inch])
            bets_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ]))
            content.append(bets_table)
        
        # Build PDF
        doc.build(content)
        buffer.seek(0)
        
        return send_file(
            io.BytesIO(buffer.read()),
            as_attachment=True,
            download_name=f'relatorio_apostas_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate-word-report', methods=['POST'])
def generate_word_report():
    """Gera relatório Word das apostas"""
    try:
        # Obter dados do banco de dados
        games = Game.query.all()
        bets = Bet.query.all()
        
        # Converter para dicionários 
        games_data = [game.to_dict() for game in games]
        bets_data = [bet.to_dict() for bet in bets]
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('🏆 MASTER LEAGUE - Relatório de Apostas', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph('')  # Space
        
        # Summary
        doc.add_heading('Resumo Geral', level=1)
        
        total_bets = len(bets_data)
        total_amount = sum(bet['amount'] for bet in bets_data)
        pending_bets = len([bet for bet in bets_data if bet['status'] == 'pending'])
        won_bets = len([bet for bet in bets_data if bet['status'] == 'won'])
        lost_bets = len([bet for bet in bets_data if bet['status'] == 'lost'])
        
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('Total de Apostas', str(total_bets)),
            ('Valor Total Apostado', f"€{total_amount:,.2f}".replace(',', '.')),
            ('Apostas Pendentes', str(pending_bets)),
            ('Apostas Ganhas', str(won_bets)),
            ('Apostas Perdidas', str(lost_bets)),
            ('Taxa de Vitória', f"{(won_bets/total_bets*100):.1f}%" if total_bets > 0 else "0%")
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
        
        # Detailed bets
        if bets_data:
            doc.add_heading('Detalhes das Apostas', level=1)
            
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Headers
            headers = ['Jogador', 'Jogo', 'Aposta', 'Valor (€)', 'Odd', 'Status']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
                table.cell(0, i).paragraphs[0].runs[0].bold = True
            
            # Data rows
            for bet in bets_data:
                row = table.add_row()
                
                bet_type_text = ''
                if bet['type'] == 'home':
                    bet_type_text = bet['gameDetails']['homeTeam']
                elif bet['type'] == 'away':
                    bet_type_text = bet['gameDetails']['awayTeam']
                else:
                    bet_type_text = 'Empate'
                
                status_text = {
                    'pending': 'Pendente',
                    'won': 'Ganhou',
                    'lost': 'Perdeu'
                }.get(bet['status'], 'Pendente')
                
                row.cells[0].text = bet['player']
                row.cells[1].text = bet['gameName']
                row.cells[2].text = bet_type_text
                row.cells[3].text = f"€{bet['amount']:,.2f}".replace(',', '.')
                row.cells[4].text = str(bet['odd'])
                row.cells[5].text = status_text
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'relatorio_apostas_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)